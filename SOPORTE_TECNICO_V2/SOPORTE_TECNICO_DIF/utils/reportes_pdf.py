"""
Módulo de generación de reportes - Sistema DIF
Coloca este archivo en: utils/reportes_pdf.py
"""

from datetime import datetime
import json


class ReportesManager:

    @staticmethod
    def generar_reporte_texto(stats, tickets):
        lineas = [
            "=" * 60,
            "ITIL HELPDESK - REPORTE DE TICKETS",
            "Sistema DIF El Marques - Soporte Tecnico",
            "=" * 60,
            f"Fecha: {datetime.now().strftime('%d/%m/%Y %H:%M:%S')}",
            "",
            "ESTADISTICAS GENERALES",
            "-" * 60,
            f"Total de tickets: {stats.get('total', 0)}",
            "",
            "Tickets por estado:",
        ]
        for e in stats.get('por_estado', []):
            lineas.append(f"  - {e['nombre']}: {e['cantidad']}")
        lineas.append("")
        lineas.append("Tickets por prioridad:")
        for p in stats.get('por_prioridad', []):
            lineas.append(f"  - {p['nombre']}: {p['cantidad']}")
        lineas.append("")
        lineas.append("Tipos de problema mas frecuentes:")
        for t in stats.get('por_tipo', []):
            lineas.append(f"  - {t['nombre']}: {t['cantidad']}")
        lineas.append("")
        lineas.append("LISTA DE TICKETS (ultimos 20)")
        lineas.append("-" * 60)
        for ticket in tickets[:20]:
            lineas.append(f"Folio:   {ticket.get('folio', '')}")
            lineas.append(f"Titulo:  {ticket.get('titulo', '')}")
            lineas.append(f"Estado:  {ticket.get('estado', '')} | Prioridad: {ticket.get('prioridad', '')}")
            lineas.append(f"Usuario: {ticket.get('nombre_usuario', '')} - {ticket.get('area_usuario', '')}")
            if ticket.get('nombre_tecnico'):
                lineas.append(f"Tecnico: {ticket['nombre_tecnico']}")
            lineas.append("")
        lineas += ["=" * 60, "Fin del reporte"]
        return "\n".join(lineas)

    @staticmethod
    def generar_reporte_json(stats, tickets):
        return json.dumps({
            'fecha_generacion': datetime.now().isoformat(),
            'sistema': 'ITIL Helpdesk - Sistema DIF El Marques',
            'estadisticas': stats,
            'tickets': tickets
        }, indent=2, ensure_ascii=False, default=str)

    @staticmethod
    def generar_reporte_csv(tickets):
        filas = ["Folio,Titulo,Usuario,Area,Estado,Prioridad,Tecnico,Fecha Creacion"]
        for t in tickets:
            fila = [
                str(t.get('folio', '')),
                str(t.get('titulo', '')).replace(',', ';'),
                str(t.get('nombre_usuario', '')).replace(',', ';'),
                str(t.get('area_usuario', '')).replace(',', ';'),
                str(t.get('estado', '')),
                str(t.get('prioridad', '')),
                str(t.get('nombre_tecnico', 'Sin asignar')).replace(',', ';'),
                str(t.get('fecha_creacion', ''))[:16],
            ]
            filas.append(','.join(fila))
        return "\n".join(filas)

    @staticmethod
    def generar_reporte_pdf(stats, tickets):
        """Genera HTML profesional de reporte general para convertir a PDF."""
        from datetime import datetime
        fecha = datetime.now().strftime('%d/%m/%Y %H:%M')
        total = stats.get('total', 0)

        color_est  = {'Abierto':'#3b82f6','En Proceso':'#f59e0b','Resuelto':'#22c55e','Cerrado':'#94a3b8'}
        color_prio = {'Baja':'#22c55e','Media':'#f59e0b','Alta':'#f97316','Crítica':'#ef4444','Critica':'#ef4444'}

        # KPI cards
        kpi_html = ''
        for e in stats.get('por_estado', []):
            c = color_est.get(e['nombre'], '#0d9488')
            kpi_html += f'''<div class="kpi">
              <div class="kpi-num" style="color:{c}">{e["cantidad"]}</div>
              <div class="kpi-lbl">{e["nombre"]}</div>
            </div>'''

        # Prioridad resumen
        prio_html = ''
        for p in stats.get('por_prioridad', []):
            cp = color_prio.get(p['nombre'], '#94a3b8')
            prio_html += f'''<div class="kpi">
              <div class="kpi-num" style="color:{cp}">{p["cantidad"]}</div>
              <div class="kpi-lbl">{p["nombre"]}</div>
            </div>'''

        # Filas de la tabla
        rows_html = ''
        for i, t in enumerate(tickets):
            bg  = '#f8fafc' if i % 2 == 0 else '#fff'
            cp  = color_prio.get(t.get('prioridad',''), '#94a3b8')
            ce  = color_est.get(t.get('estado',''), '#94a3b8')
            rows_html += f"""<tr style="background:{bg}">
              <td style="font-family:monospace;font-size:9px;font-weight:700;color:#0f766e;white-space:nowrap">{t.get('folio','')}</td>
              <td style="font-size:10px;font-weight:600">{str(t.get('titulo',''))[:52]}</td>
              <td style="font-size:10px">{t.get('nombre_usuario','')}</td>
              <td style="font-size:10px">{t.get('area_usuario','')}</td>
              <td><span style="background:{ce};color:#fff;padding:2px 8px;border-radius:20px;font-size:8.5px;font-weight:700;white-space:nowrap">{t.get('estado','')}</span></td>
              <td><span style="background:{cp};color:#fff;padding:2px 8px;border-radius:20px;font-size:8.5px;font-weight:700;white-space:nowrap">{t.get('prioridad','')}</span></td>
              <td style="font-size:9.5px;color:#475569">{t.get('nombre_tecnico','') or '—'}</td>
              <td style="font-size:9px;color:#64748b;white-space:nowrap">{str(t.get('fecha_creacion',''))[:16]}</td>
            </tr>"""

        B64_LOGO = "iVBORw0KGgoAAAANSUhEUgAAAFoAAABaCAYAAAA4qEECAAABCGlDQ1BJQ0MgUHJvZmlsZQAAeJxjYGA8wQAELAYMDLl5JUVB7k4KEZFRCuwPGBiBEAwSk4sLGHADoKpv1yBqL+viUYcLcKakFicD6Q9ArFIEtBxopAiQLZIOYWuA2EkQtg2IXV5SUAJkB4DYRSFBzkB2CpCtkY7ETkJiJxcUgdT3ANk2uTmlyQh3M/Ck5oUGA2kOIJZhKGYIYnBncAL5H6IkfxEDg8VXBgbmCQixpJkMDNtbGRgkbiHEVBYwMPC3MDBsO48QQ4RJQWJRIliIBYiZ0tIYGD4tZ2DgjWRgEL7AwMAVDQsIHG5TALvNnSEfCNMZchhSgSKeDHkMyQx6QJYRgwGDIYMZAKbWPz9HbOBQAAAu3klEQVR4nMWdd5xU1fn/3+fcO31md3bZpYgUpVmIGktUsFdsKEpHYsREoxIbdpOY5Btj11iQIJZYEAELtogtKqhYsBcQrDRl2WVnp8/ce8/5/XFnZmd3Z5cFzev3vF6zuzNzz7nP+dxznnrOs+KFl9/TAEJqiqS1RghR+rv1CwG4nwsh0Cig7FolqESVPi32KoRw71exZeek278Tlflsw0fhXuXjc79QFa+FduMHhJYd+Ci/pthOCEEun2dg/54M32UQZncG1Rm5fYqy920H157Jtm27hrazgW4Ltb+XEKLsM02lO2ztfbc0ntLjKT7p4g3K/27XZWlmbOlmXd68rP8tMflzUTnPP8cDBEAoVxoIVXFlFKnDjN6ameReUs686jZo5csYOi5097viqvlp9LOB+hPIFEK3WTxdMiU04BRG3w3mC7JXFK7XZXJUadUGXa1pi7aW7v2k29ZxHPeWQiAEyNLFEhBoLQr86dbf5X1BJytRlL7vfMxl+HQxa93rC98X+CxOPLOyhPrfUacPUlBRQSmlEAgi4QhokFKQy+XJZnMYRhcAdYsXRVei8OekbVKGXYHVhmRB7pcadt5nJUWqlCIQCIGWvPzyEl5//TXq6usZc+JJ7DBwIIlkC46yMKRR4eZd8F02KwVFEfjzU/mQfpLV0YF02/WxNfK1/OEppfB6vQT8IVasWMUDDz7MG2+/B0KgHIdXX3+LqRPHc+xxRxMJR0gmk2gE0vhpAr07M3tbrSHxwivvbLFFydbt5hLrTIl13l6DFmilkKYkFA6zceMmFi54imeff5FsPkc4Uo1tWUhp4jgOqWQTu+60E1NPncL++++LY1tkM2mklAUwugBdFH+J4t3b8OL6Bz+NhBBYls2Afr3Ydecduw90iY1uPslKoEopUUp1uE4pDVoRCkdIpXMsfuElFjz2BD/+sJmqqhpS6SRWNkNVdYhc1iaXtQjXVGFbOWw7z4EjRzB16kSGDtmRTDqOnXcwDE+3xgOtQLv8tnN+tpHaA90t0bEtxnulNkqpDl6Z42gC/hBIWPLmMh6Zu5CVK7/BHw7h8weIbd7IAQfsw7TfjGOvPXeneXOMp595iQcefoxkS4qqmlqWvPU277//PieMPoZTTh5Nr7p6UskUSimk7Kgw2/P2P1ODZR13mNHbIoMq2c6V3NLyzx3Hxuv1EfBHWLHiSx56ZB5vLluGxxtEGl5SLc0MG7Yjl844k1MnjcHr9bXp/7PPV3DdTXey4PHncLSXUDBMMh6jb9/tmDz+JI4+6jB8Xh/JVAqgIuAlntq/+7lmdN5mQP+i6CjGOkpYlIFSiGW0flCUfdqVbxXaQCuYgnIHRrhxAeUghCjJ4fkLF/Gf518in7fxh8K0xDbRu2cN55w1jd+fdRr1PaoBeP+jT3j9tbeo71nH6OOPorqqCoCXXlnKtdfP5NXX3yEQCCM9BulUjOG77Myvp0xm/31/ieNYZLI5pDTRSMRWyOAtQV5ZE2iEgHzeZsCAMqAre3MCUG0VQwHoohPSGROVgFbKQQsIhyOkkhn+s/glFj72JBsbNhOJ1JCMx/F4DaZOHs0lF53F4ME7APDNd+u46eZZzJ33JPFYGoRmt92Hctkl05k84UQAbCvPvAVPccONs/js868IRmtwbIWysxw4ch+mnjqJYUOHkckkyeVzmLJz+d0Rsq5JVpQAFYB+8ZXlunhhMdhSKc4hRKu92ertbYGpsrhJIOADCUvfeJe5jyxgxaqvCIZDOHmLXCbD0UcdyFVXnM8B++8DQHOsmTn3zuWfd/ybH9ZvIlxdjWl60GhSqSS2bXPc0Qdw1eV/YP/93DZNm1uY9a8HuGPWfTQ0NFMdrSWZSRL0+Tju2KMZP/YkevWqI5VI4DgKw5AFZSy22dWvbCi4QOfyFgP692L4LoM6At22QecROdkNoJXj4PP58Pt8fPb5Ch5+ZCFvvf0OpjcIpkGmpZk99hjGZTPOYcL4kxCA4zjMW/AU1990F599tBJ/dQ2+gBdl5QsuOEhpgDBJNDcT9BtMPfVkZlz0ewbtOBCAL1d9xc23zuLhR54h5wiCoRCpeAvb9e7JhHGnMOrogwgGgySTSYSQSGl0SyeVK/mu9VIXQCvlLvMOAJaWvitCRLvv2lsRAo2jHAwpCIZDrN/QwMKFi1i8+CWytsIbCJGKbabf9r04f/pv+e0Zk6iOhAB49fU3ufbGmbz08huYniDhUNgVOaiKa9gwTJSjiMea6dW3nvPOnspZv5tCj9oeALy29B2uvfZ2XvzvW/i8QTxek3Q6zi7DhjJlygRG7r8PoEins66yLKxaUUHylkSg1m4Mp6BcBW0nZnmcPZfPM3BAT4bvXGFGt7+FKutElgGvC15JOdBKWwgEoVCERDLNM8//h8cWLqKpKUk4UkUi2Uww4OW3p01gxgVnsv322wHwxcpV3HjjLOYtXIRlaaqqalG6aApu2UU2TA+5XI5MopGddh3KpRedw6mTT8FjmmilmP/YU9xw40w+/PhLgtVRcDRWPsv+I/Zh6uRJDN9lKNlMmlwuizRMyqLHHYAuBvpLAf6tBbrUYdnfxTiFKADcIZRZFtlyHMeVw8LD60uWMfeRBXz19deEQmHyuTx2Psfo4w/nqsvPY6+9dgOgoaGJ2++8n3/NeZCmpgRV0VqMgpvt4JQYFkK4Hh+tYXqlyvWIQAiJMA2S6TROpoVDD96PKy87nyMOPxCAlnice+5/lH/edg/r1v1AVbSGTDaP12Nw/KjDGDfuFLbrXV9w541S5E2jXYuvDOj2eP0koFsDMKLDd61At0bYqqqqWL1qDXPuvZ9l776P1xfEQJCMb2LvX+3GVZdfyEknHAVAJpvlwbkLueWWOaxa+R3BaA0+r4lj2y6jQiOFgRCCfN4il8/hWDZoV4QIj4nX48Xn82EYAsdRBX4lUkikIWhpacFjKMaOPZbLL5nO8F2GAfD99+u56bbZ3P/AI2QyNuFIDYlknPraaqZMHs+JJxyLsh0s20ZK2UEm/2xAd0ntDHl3PivCwQjPPrOYf919L4mMRTAUJrG5if4DtuOi83/DtNMnEwmFAXj2Py9z/Y138cZby/EFggT8QRzlILSDRiOkicBDIpVA5dL06t2DHQdsT9++2xGNRslms2zc2MjadZv4ds0actkswXAVPp8H27EQWhYCTAZaKeItm+lRW81Zv5vCedOn0atnPQDvLP+Q6268naeffg3TE8Tn85JIxDjooP2YccEfiEaCZPJ5hDAAVTRqO590RQur8FPIYs6w17YCXfb0lCIQDnLvfXN58IH5hCNVZPMZfIbmzGmTueDC39Ovbx8A3v/wE6674U4WPbUYLXxURapR2sFxXAcGoTEML5l0mryVZMS+uzNp4ikcP+pwBgzo14GVbDbFW8s+YtGiF1jwxH/YuLGZqmgNyBxKaYQ2ADBNEyuvSMY2M2TYAGZcdA6nnToGv98HaJ54cjHXXHMLH3yyipoePWmONzNsyECu/dsfqa2Jks3nXP60bGt1FCFpZ6m1Au1aHd0G2rUzCxaH0KU7KMehKlLF/Q/MZc69D1FT14vmpk3su89w/nnTX9lv370AWLN2HTffPpv7719IIpGnKhpFC9COotzlMTyCWFMLQ4YM4E9/ns7U8aeUhrPiy9WsXLmazZtj+P0++g/ox647D6G2pgaAr7/9hutuns2D9y9Cenz4/R4cxy71rSR4DINMOkMunWbkgXty1eXnccyRhwKQSCS59oY7uPGW2QTCtSTTCXYdNoib//FXTK+BrRTtM+udmd3tgR7Qrxe/2KWbM1oX8kxCuOknpRSRSJhXX1vG1Vf/g2i0B02bG5gyaTSz7riOSCREJptl9uyHueX22axd00CwugavKd2UVDsyDINYUyPjxo3ijtv+Qa/6Oho2NXLPvQ/z2JOL+fqbdSQTqdID9/kD1NXXc9DIPTjjN+M5/DBX4S16+nmmn/dnGhoThMLB1nsVdIkUBkIatMQTmMJm7JhRXHn5eQzfdScA5j36JGeedTGGP0pLPMbxow7nssvOJ5VKIkVrcqGr7REdZnS/dnZ0Z0+qFDoEVCG/6DVMEqk008+/lM3NGTLJOFMnncB9996GIeGjTz/n7D9cxdtvLCcQieL3B7AdG3QZyNq1IgyPpLmpiT+cezq33/oXAO759zz+fs3tfP/tGsxAGJ8viMdsNbkcrcnlcuRTSTw+D2NGH8N111zCDjv059PPvuDkcWexZv1GgsEgtqMQ7cw1KQ1A09LcRLSmir//7RLOPevXADw6bxFTf3sBgep60i3N/OVPl3LooQeSTCbcdroY66kceOsoOnozfOcdKxiM7ciVn63JWDe9FGDRU8/xw4YGbCvHL/fciZl3Xoch4c0332PUMRN5+52PqanvhcdjYts512IoKRSBlg7SYxJraubUU0dz+61/IZe1OOvsS/jd7y7kx00t1NT3IRQMYkiF49ilF8rB7zWJ9uhBIBhmwcKnOfjwk3nl5df5xfBdWDD/bqrDfqycQko3ZlNU3aBRykYph5oedeRtmH7Opfzxj9eBhomTTuLcc39DorkRjyfAwgVPkkln3X4Kzsq25L66nd3Uwn2CXtNDQ0MTr/z3dfz+EFLnuPn6PxEOBfh85UrGTzyLzTGbaLQGy7JKYqfDA5SCVCrB8F2Gcuetf8e28/z2zAu5+1/zqOq5PX6fF9u2CjHsCvxo13bXWhOt78WPjWnGTjyD1159g1/utjM33/QXMulUxVBBkWzbxjAl4Z59uOaau5g56wEA/nj5BQzaoS8AK1d/zfL3PyQYCKKc1mhmZy57aSq18zC3CHSbmJTSBP1BPv74UzZu3Ew2m+PYow/lkAP3J5/Pc+EFf2HDxgaCkQB2vqMsLidDe9FWjr/+7SKqq6u5/qbZPPzwAqJ9e6Ity826dJMcK0Mw5CeRM5l25kWsX7eOUyeNYdQx+9HSEu8y24LSSAWBHrX86f+uZ8WqVdTVRjnzjClkM3E0giVL30RpED8hiys7RunazYDCRha0KEWi313+AUqD1g4TJo5Ba3h80WJefuUtoj3qsO08QlYOvABIQ5JMpBmx/96MPv4ovlj5NTfePIdgbU8cKw9aYRais7YhEFog0a6OEAX5KARFI0gLcGybSKSKb7/5gb9eMxMhBBed/zvMUvqsM5DcXKXPa9DclOT2me6sPvGEo6mpjYCQrFi5mpaWJKZplsRGdzYKldspHWZ0V1vBpCFJZ9KsWbserTR9etczcsQ+CAFz5z0BhoFWAonRJQNSCGwrx/hxJ2JIwR0z76GlOYbP4wNloJFkTZAaPI5CGRrHNDGkCYaJNiRSK5QERxhoTAQCZecI19Ty6MJnWbl6JYccdCB77L4r6XQKIbpevNp2CISqePHF12hqjjFk8A7sNnwXbNuhuTnGDxs2Yno8hb0gW0+yszBoh7ya1hjSIJlI0RJrwXEUgwf1p2+fXmxqbOSTT1fg8we3zIgAy3aIVPs58oiDSKXTvPjS6/jDVSVzTAI+B/KGK8u1lSefaMFqacGJxcinE2g0hpBIDV6nGOwBj8ck0dzCvHnPYBgGBx+8L3Y202UqC9zgmdfrYd36H/nkky+QUjBk8CBQimw2y/r16zBNs10it+uZXT5pK969s8aGYRKPJ0im0qChrs51GL5fu46mzQk3ML8FMSaFIJ/N079/HwYP7s8HH37MmrUb8fr8bR5u3tTYhiCbzGNs15voSUdSf94UekwbS+hXe2BrcNI58l5F3nBKS1QphfQHePXV99DAyJF7IwyzQ/a9s3Hn84p1GxoA6NOrrsRTOp1GCtk24boV2YKOWfCyFFUxFCpwn7ibQrdwHAVaEQkFC0xksWwHn69g+nSaGnOFquXk6d+3F1II1ny31tX+CCyhXQtVCDy2Bseidtpoor8+CdWzFssQSAV1mSy5Dz5l3a1zkSu/xuv3YpU5Ex6Phw0bNoKTZ6fBgwhFQthKYQiJa+pVBkgIwNE0N7cAEAwGAHdW2rabiitGD7tOYot2vzuZ0aVtXO0Ak+gy0aBLGw2V475396107TMJDNCSqmgtm5XivVQOKQyUsCkaRBLIWhY9LzyVmivOIFUdJpVOY7ekyMWTNCkbe8Qv2WHmlfiHDyGfsUphVASYOk/MULwVi2NEggQCPrSWdAZwe5BUYYyuWC+EgbVyQS4DuOusTHvsuqKyPdNtNyC276s487sYiHAfnpQmOm8RMz2cvvEbnjtwONuNOYxMOoEUEhOJlcoQPXI/Qr8+hU2JHDg5DCkRhvvyYpJL5Ij1qmG7P/4eFa1GWLZrMzsaAgGq/ng2Z8R+4L6GNfiQaMfuMqFcCSLZRoFWXqndFR9dAt2BqUIwvAho+VaCLZEGhCGJx+PsuH0te44bxYfpFKo6QnD4ELTj9uegcIIeqscdSxaHgGOjpIGhXfkuC5E+nzDQCQux6xBqR+1PJp9FSANl2xj19QR/MYRNlsGPdTWMG38s2UwMrR2QXVtEHak4gzsCUkxKdAfsDkCLdh0UB1fydir0KajARTsyhEE2meKAEb/grbcXs/8hI7CTOYR0cCyNocGWBo5l499uO7yDB5F1MiiPhVdBi2PTknGI5WxSKCxTIA0HR0F4n93AkDjCNfkEgJ3DMPJoy+aWm//KQw/cgXJyhVhF11SeuSllhItiU7e9bssixCVZfnGpUdmrffag4LuU3hfZEZ242pSuEWjH5rJLzqFXn57kMxnXHka6wlAo9ylaChmNkA1DKGeg8iYtdo6DpZ+b+/fiyp517JCFWD6DYRvk0dg9qzB8PoTjuDxpQHsQGhzcLMmpE8ewy87DSBc2QnYBc4c9GiDd1JZ0HaU233bTW+xWrKPtUysDs8u4X1tSaDxeL8FQuCDzpbs1p5CX04DX0SAcNDZeBWkB/aTk7r6D+X3f7VifylLn8XH3oMFc1aMPQmscLZFKoYTG1GBoUKLITWv6yVYKX7ttZZ1ShaGUxr+NXvhPO5UlywLhW1qRpcxEQbEWEwjF1aEFGoH2SkRDDCPukKvRmELw30wTLzc20iQ8iEaHPUIehkd7YGobDIH+IQbJPHaVH+yixVTkqTy7v3UodZys3beb25NZeUtTZ3fu5G0J5a7NHdvJo4Rrk2vHxjI0hjKwC/FdBwfD68Ve10D60y+JHrUP65uSfGolqfFVEZIGQis+VTYfb2zGKyU+J03Tc69hoLClicRBalVIZUlsNI5jYRoe8ra1VeBUGnD3oW7btiQ6uqs9y5VJ+WrSW1AyUkoc2+HBh+cjhaAqFMa0NGgTWyiEFu5ePeV6X7EHF+FPZcmHfPRUErCxtY0lLILaISQlMuhD4iG48xC0pdF5C6RECQuFgxI2Hq0wDQ9L31rOihWr8fv9nU+qwoGj1ox3q5IXhehd24RWpyBR9CuKfcn2WrM94F2BvzULyXEcwlV1PPTQk0yZeCavv/shAdMkp/NIIVttcKWRQR/JDz6i+R93U6MVKhTEFgZCCVAS22PiiQYRX3xHfmMzvsun0XvGNLR0sOMJhPS4QSalaXE0N992D1NOPQfLdq2f/x/H4UoyusPWrm7M7tbjbKIbqLvmUTBUyyPznyaSaGbAff/A8Zikf2zCkVCMGmut8YYDbH7yFTLrN1Fz1slEdh6GDgfBsVENm8g99yw/3Psk/kF9qf2/s/FdPom+x+9N6tZHiS/7BJGx6D2sP1889SpPz/gzgeo6N56iLLo/RcpXb/fMuDYjLmvTqgwrAtt1x0WcO2OiqJRc2aYBheNY1PTqg/P+apqvnoUZCNH01FICPn/hqJty2ymBJ+gj/e7HJD9ZgXfIQCK9a7EzGRLfbECsa0RXQe6Tz9h4+tVUnXwENWOPwBy5Jy3PvM76a++kdp9daV70OpGaHmB40Y7VahpXGm17CMrjNqJ7k68MnTbvTE1ZgKTs69akbEf7uNzzpuz69vwXGZeFn0oItNDklMbQDpsefgoU+AJhbNOLoRS2ITCVRgmBJUBGghgOiM++pfnTlYBEeIN4ImEsLIxQkFw6Q8OcBSTmPod2BGZ1NdmlH7H2lfcw/T4Mj8cFuZypCuSu0PJAgm5tUrCUujenRdnLpU7Nu5966L3YytDuzHcMd6uCqSRKgxIGgWAEbYBWDlLZpH0Cr+0yKlAYGijEmu2wiYHHVZoKhONgCtBY+ISAcBjt2GBoHCnw+n0Y/gA5dGGl/P8ls1spGSHa5MsqQ18h4FIwhyxT4CiFsBV5mcfnbjRAKBPLtlEC/I6NX7lh0JwpXKVY6FejQDkI5R6LUCakpSZgCfKGgyENTEfjGBKlFUZOkZUKJTXSkHi2kPFpTx1WZvH3T1Cinc7ock+wvevdKUdlVLpUCvKZLH0nHEvguEPwppKsvWMuuS+/I+fk6H3aaOSRv0Kkc+Tmv0pg1H7o2mqwNba3YBo5AttxMIJezOWfsXbOQob++VzSA+rxf/MDa6+dg2M7yLRFYLch9DhvCsLnh2Uf8v2s+UivWZT8XZIoDbC9C/7TqQPQ5XvLCtZgSVQr3Dxe6wWtselKJ5l0IUBs5DVy0PbYB/wSkYxj3vc0We2glY0xeCBqxK+QySTm21/gO3hPkv164k3l0dkMUguE9ODxeZHRIMIjUA8+idrzF9hDt4Oqb9EekJZEqTxEq9Ej98D2B/FsbgbHQUt3hWyJ3CGUHV9uH6DojiXWtkHpr5/3iHJnJEDlLWQqjZPKoGy79IWTy6NTaUQ+S7qhgdw1cyBahd0rimfMYWifH/HNd6Se+S/SH8D+dgOmI3FSGXQqjZPOlLDRwgVUp7I4Dshs7iexrUsRqiK3Wyfpy59LG6C11m03nHSQSa0BoA4M6cpHmHV5N4ZESFlICRVmvBRoKcHrwUpnaVq0FCEhsttQeo45iow/CD828eOcBYTwYEmB0StacKxku+B8YXSGQBiGe/zhpzgnHbZitI6pFZGuqPUQktn24y7uKUpR2dLdiiLNdTUrBwJdy6ggjtzNHqVdPLrYn9auiDAkoXCQvE+go34EBl5HoE0TX7Aa7ffiyTvu9lDt1uzQJTu9mJAo1ggpapetpVYbTou2JtrWP7PWBrLc4tvSmW+hBVoUTSXRZlG5DFZSItIFWgu0LIQwdUHxaIEULhPuJhmwtYN0FNp2UChsw3aDUMrBsB1QNkq4qk1qcKQLckl36OLjK+qarVjuBT7LomVteqjkK3ToouxVrKLjolCBKh2r+F/Ttri4nVG57dDqbBRf3eDlZ+GiLXWpDNt6q7oU36XMMmnboP2sLs4tXZLVQrifVZppxetEu09lSfYXvi3Ng47gCSGQWrgZeq3RloXOW2hdOB4hBWxhM03rDO44yG2dev8zq6McblnaU1x0cVvrK7W10YuiyQVCCY0uKM/2IQqpFQqBqZW74stEhZIalAU+D6JPHTrgB1u5UiGXQ6ZyXa7WkgUrBO0X/dbN9jLzzrVg2mvXVve7LJ5XuKxcQZQv0naJSyHctFIhuC+0QiuwtACpUCikdvvUjsJwFKY2MAoAogxMR4EtsFCIkrzTaCFQWOCAdGxsKfAKB9C4BzYEOpFG7fULBsy7FbAQeQFhL847H7DuqrvweDwVRFW7mI4GCt6o1hLXV+1oVXVsSetjr2R1FEH+OWRl0QLQQmMJG680McJhApZDCgOvLcgJB5/XA8EAnryB7TFRSEylkdLAFw4gqvwEDB8/FgegNcKGoM+HN+zD8AfRys2mCAUeqaj1BdhseNGGgB4ScMhpMMN+xPoff1Z90F0ywUW/3H7u2vooe99+Ypet76KjKB0HM+Cl+ZklyC/XIvI5Mt9vQHk9SKBx/mI8b32OcGz4aBVe08QyJfL7DfxwyS1ovxf9QyPSMFxbXUp0Ls+G6+5BRqtx4kl8Wdf7M7wekqu/Zc1l12M7Gl0IJkkEjtAYhoHaGMOURmU9oXWbcKh7Gku5+Uytt8k2KMLXJvBfDnIHgMtuVNnIqaTcQCjtnohasZr051/g0waG14MyTQwhSXy0Et5ZCYaNGfCCaSKkgZWJk372NbRtQ8BHwB/GsW2UdtBKkVnyAY5tYUmFxwyhBJhC4mloYePCV1AovIbp7sUQ4BEeHDuP8noI+oPdi30Ulf/PEPszaSOHO4tQadxiIk5roVdJSaG5qSjZTny5mwk1HhQOvogPHwHyeYXXI9HKDT/0rIsiTJOslSZvKbQwcPIOQkqq6qIEfCaxWJpUMkl1dTU+v4NtaTY1pAlFI1T5wCu9mB7I5zSNjc307l+P12fS3JDG9JhowyYRs4nWBTEMP/FYsnLNPAFoQdFTcD1XBUIipYtB+1XQ2SMQJQyKQG8rqdYn7W4uqTxDDMOgpbGJs6+YTs/6Hlw+42LOveJKDEOyYe16rrxsOvFEnEcWPEsoFMYwJH/983VU1dbwwL230LsuiqXhoguu4sCDDuC0qWPY3JLi448/48IL/sRfr7qa40YdQDKZYu6CZ/nm62/481UX4DFMLrv8WkaPHsV7yz/lnXeX8eDcuzl10nSaG2MYhreTgekSQLLkzv20EClUVIa6rOMyC6Mot2Tr+3zeDQ5JQ3a5uLTjUFsTpaa6CseOEfL7CYS8SHqxePGrXDzjbwhTMnPWDSSTLWhHIYWfPn16MXn86Rx8+EiuvOJ8vv72O55/5VVuue5Gnn7uGc6Zfga11WHmPriQ6278F4bfy1tLn+G+e//NzNnz8Pv9nHnuFHZK9ue000/hrpn3s3Lll1T3qGs9sVvOpwYk+P3uQ3Bsp/S5ayRsG8jQzkjcUsZbKUXAH8BrekAIYrEWAHrW1xEM+QvMl7cokOFh8+YYRx8xkgfnPcHkiSeQiKeJJzKMOuZQHl04iyG7DCORTKG1wjQDaCnIWSmOPuYoDjnkEL5ds5acZZHOaNav28QdM+9nv5G/JJFMMX7iSTy6YDbDdh3GnTPv5YrLZ3Dnnf/Aa5hsXN/EpRefyzvvvst9s+8h2rNXm1O15aTRCENQX98DraGxcXMBF43P181dTp2Qu+ALHpamaKBLhDDKAkUCLQwsxyEarSFaHUUYknUbGshksuw4sB/Dhgwik8sWoiflQCtQCq/Hyyeffc6su2bxznvv4zENAj4/y5d/zLXX3sGqlV/h9/pJJlLY9joSiRjCgaOPPIAVX67myiv+TihYRSaVABIM2rE/yUQOj8fPC6+8xt+vuYOG9Zu5f85D7LffMQweuD1XX30BqXSKF15+jRH77kv/HQaSjLsHM0UbU0mgDI2lLGqqI+y281CEgHVrN4IwMSTU19SgVasjWtz42V1qFybtWJa4tAEE0EoRDgXp0aOa79cZrP7qWz7++Av2229Pjh11KMuWLMeMRLGwAIdSQEYrqqPVbI6lWLb0JY47/kSqq0M4do6BA/qz87AhBCMRTK/g2KMPIR67go8+WYVp+vjjX2/mg3feA0z8Pj9HHTYCr7iWab+ZxEknncalM86jjyfE7rsOpk+vFvodfzjxWCPxeJxUMsXQYUN55tnXCIcljy68n6OOnOAeyi8N05XHhvCQTMY55MgDGDpkEPF4nE8++wzTY1IdCTNw4EDyVn4rM+FlQLct00Ob/R0dTD6lCfg8DN91Zz748HMymRyLnn2e/fbbkzNOm8g99zxEw6YkXn8ARzsIXTiEEwzx8stLiEbDeMI7sOSN9zAwiLU0EwpH2GfEHjjLP2XuvCdoajyY7QcMZMXqtdx2572kM5JgfT/sXJaFC5/liCP2xRusZuyEaSx//xMWPP40++zzC/bcdx/eeusd1v+4jpNPPIFVX63h5ltnc+SRh9HQmOSNJS9heEIM3HFHVq36ut2OpUKIVWU5+8ypCCF4/c13+fKr7/B4/fTdrg919T3I27nSKi/i1F0l2fYseIWTRuWhQaEUgYCPFau/4eJLrsZGUxMxeHvJ82zfrw/zFjzJ5MlnU1XTB20AhVNWQkpy6Sy2svFGAlipLFpJpOFgJXOgJfjA6/GTT6dBORjBIA4KvyeMx+OghSSXymNlE6AkMhgkWBUmHW9BpS1AI3wBEAqdTYOEYFVPsvkE0pAEfRHiTU0Eqv2YHh/lCVDT56d5wzqmnTmROXfdhJBw/Jjf8MJL74K2+d20qUw9dTyxeDOm9LQKnS6ALhZGGVCq8V8s8asLQfmyi8viZYXGBplMhmFDB7HHHrvxznvvsWFjlj/97Ubun3MLk8aPYd2aDVx6+T/wVdUS8HncY8RK4Q8GAIFyLDyBEK7qUYRDVe7nhVNTwVDIva9yBaJWCq0N0BAIBQhFQqWYh2M7hENVyIgbU3AKR4hldQQQ7jY0bxVohVIO1fU9UMouOIAaKSXSMGn+cQNHHHUgN1/3F6QULHjsGRYvXkogXE3Qqzj4oJHksnkMYW5RLnc2v7ey0rVAu0eGmDB+NNKQhKvq+PdDjzPzXw8BcMnF5/LAv2+jOiiINTUjhEQYJk5hsBQ2hrtRT4ltO9i2jVIKpRSO46Acx33vtK2r5J6Ocg/eF/dqKOW2t22n1K/jqNKZRV04Sy4Ehfu7qTPp8ZDLWTQ3NDBp8mjmz7uLaDTCqlXfceFFVxMIRkglN3P88cfQv9/2ZHPZtiudrbOtuwTajVyKkg2thUZIg2wmw+6778q4k0eTiMWIRHtywYyrmXPfXAB+fepYlry6iClTTiSXiZGIxzCkscVDlW3uXbbZsjwIVGkT5tYoKMN0zxy2NG6kb58oc+6+kYf/fSe1NbWsWv0NY8adQVNzBsvOsdOwQYwfdzKpdKLgvLSu8vK/dbvPK9FWzWgtHLRQSOkhncxw2tQJjByxN4mWZsKRKGef8ycuvez/iCeSDBuyIw/f/08WP/MwRx46gnismXQ6jWmaneYXuyJpGNiWhW3bhXobW0eG4d43tnkzAY/k8svP441XF/HbaROQUvDSi29w5DGTWP3tWgy/SSDk4+IZFxAIeHGUXZhr2x7zkEJLRKFIScWYanE2CVdslOIAWiOU4spLL2Sv3Xcl1hyjqqaWG2+5j0OOHMvjTz0HwCEHj+Q/zzzIvIdv4xfDB9PcuJGcncfwmAUvtOu5UCy6HW/YyNFHjORXew+nZVMDshTr7Sw24/YrpUAaXuLxBLlcgl//egxLXp7PtX+7lN6961m7bh1/mHEVJ54yjabNObweH4ZW/OXKy9lp0CAyqRxSmHRMPXTCb9mrzecvVSiZ2QHsAhWONCKLLqnK4/H6yGYdbrntDl5+dSmRUA35fJZ8PslxxxzGpRdP58ARewNu7aJ77p/LLXfczbrvNxGu6oHpMQrVaSqXz1FKEY1WEw0J/n3fnWzc1MRFF11JLGWRTuUQ0ktr4ZMCzwKEkJjSJJlMYtsZjjx8BJddfA6HHXKAy0syxb33P8qtt89izZr1RGv6EE/G6dWzlisvu4Q9dtuVZCKGYXjYmv97VG6RlFsd3QK6mGlRAhBlGWdcBWRIA9Pj5annXmDevMdo3BQjGAqTTLbg90lOnXQiF11wFsOGDgbg+3Xr+edt9/LvBx4lFk9RVd3DVVbt6i1JKcnlcgwc2I8nF85m6I79EIZkybIPmTjhTBLJDIbpc0OhxZyGAGl6yWXSZJIt7Lb7rlx2ybmMH3c8ppSgFfMff5YbbrqTDz74gmCkBoFDNpvlgAP34+yzzmS7Xj1JJuMYpiycRJCtoeEt2M7tgR44oCe77jwI8WKhfnTxy60C2o0gorUDShGKVPHDxk0sWPikW4s0ZxEIRGiJN1NXV805Z01l+tmnUV/n1g798KPPuO7GmTy56CWUFkSqqtBatTkgbxomzRvXc8HFv+PWG/8PgJGHjmbZW58TranCLj0cgWGa2JZFsiXOgAF9mH7uafx22iSi1dUALFn6NtddP4vnX34dr8+H1+slnUoxZNBApkyewMEHHYBt2+RzeTfRUBBBbYTGtgLdWqh7C8ujTOu3P4pB4WEUq5z7AwFWfbmKeQseZ8kby0D4kB5JumUzQ4cO5tIZZzJ10sl4C4GaxS++xrU33MWSpW/j9QUIhNwtuBoHKUwSsRgXXfR7pIC6+nr+++pSXnxpKZGqiLt7VBooLYjHNhOtCXL6byZzwfQz6L+9W3Pvyy9Xc8MtdzN3/rPkLZtQKEw6GWe7nnWcfMqJHDXqCCLhMOlUEkGFQrhdwbIFoEuiY0vl2IqzuTMvqEOBPa1RWhMMBJBC8u7y95n36GN89OmneINhtCXIpJs5aOS+XHHpdEaNOgQAy7J4+JEnuOmWu/ji89UEIj3w+fzYjoVCEfR5aGxsxGOaRGtrSaazmIYHhCTV0ozXC+PGnsDFF57FbsN3BmBjwybuvOtBZs95gE2NCSLVtWTTLQRDPk44/lhOGn0CvXrVk0ynUI6DWbBmtsa2+NmALu9gi2WByq7Ryi01EQyFsCyH/762hEcWPMH3360nGIqQyWYRKsdJJxzGpZecy9577QFAY1Mzs+9+kLv+9SAb1jUSrqlBmpC37FLmWmuF6fGQSrRg5/McdcSBBUU3EoBsNsODDy/gpltns/rLDQRqa3ByOUyhOeTgkUyYOIZBgwaRzaTI5fMY5W51d8Bgy0nsbQa6UtS762LVrpOtlUIISTgUpqU5zpPPPseip5+nqTlGMBwmGY8TCvqZ9uuxnP+H00vFtr/55nv+eds9PPDQ48STCSI1dSXnJJ/Pkkm0sMeeO3PpRecwfuyJGAVnaNHTi7nhptkse3s5vqoqDC3I5jP8ao/dmTJpArvvtTt2Pks2my0Vmi2OpdJk6grInxXoIpRdlbZp9ZQ6FhAstnEcB8MwCAVDrFm3jnnzF/LiK6/hOBLTHyLdsonePes49+zT+f2Zk6nrUQvA8g8+4drrb2PRM6+iHA1K0W9AHy44bxq/nTaFqrBb5PvNZcu54eaZPPOf1zDNAD6fl3Q6wZBBOzJ54skcdOABSClIpdNuLNndMlVwfNuOSRV3RQld2Jy+9Y5KEej+/Xt2s2RmF+KibeFTQBRjE9INUlXIrDuFcvRer5fPPv+CufMW8NY7y/H4vBjSRyoWY6dddmDG+WcxZfIpBApppf+88BqPzn+M/v37c+YZp9K/n1vk+8tVX3PTLbN5ZMGT5PI2wXA1qWSKXj1qOWXsaI4/7mjCoVCn/5el8+Lk2wZ0OT4/Cej2wEFboN0lJaBwErb1evfVWkXBIRDwo7XgjTeX8cj8x1j55dcEg1VYyiGXiDFyxF5ccel0jjvuyA48bNrUyMy77mPWnIdoaExQFe1BNh3H5/dy7KijGXvKifTpXe/+pyFHYEizxO+WrIq2QLeDo5OoRXmF9CJtFdBdKcFtBdpto1DaQSCJRMKkkjkWv/Ai8x9/kvU/biISriGbyoDKMXr0EZz3h9PZbZedSabSLHpuMXfeeR9frvyecLSGvJMHZXPQiP2ZMmEcQ3caTDabJp/LYxpm666/CmBUov9vQLen9jU8Wi8uB7ZCT+XnyCmUEEKglGtahcMRfty4kScWLeKZ514lmcwQCgeJx2ME/B76b9+XVDrN2vU/EAhGMIQkm82x2247MWXSWPbZex+UY5PJpBFSlJ1BKfd6RdmhoMrU6fjKxtg5Ff9hJGRzXVTbrQBP9xnZCqCLWrv8cJLjOO6/1wsEWP3VNzwyfyGvL3kDrQ38vgB5y8IwDKQQpNJxBvTfnokTxnLYoYfg9ZikUumCZVLgrKyweJvQalf/kbOr8ZWNsSsq8tA10IVIXVfPvIPIaMdEaVBbGFCH+xb6VkrhD3gwPT4++OATHpr7KB98+EnBelHU1/fgpBOP44Tjj6EmGiGVTKK1LOw+KjO7thLoTo9ZVxjjlsi1OlorolfcqdQdkP9XJITAMAysrE0um+eXuw9n+PC/sXTp27z99jJ69uzNMcccRf/+fUgnU6TiSaQh0VLQWiquqBP+t7xuDbWZ0a22MiW0RRu52p794oyRZe9br+lsXbQVN8U6Tm0zKsWdyI5ykFISDkVKTXL5LNlcBkO2Ohztk53FNddGjOlWOd2GhUo8FgNpZXxW4L5Taj+j/x+AaV9ZIFqOrgAAAABJRU5ErkJggg=="

        return f"""<!DOCTYPE html>
<html lang="es">
<head>
<meta charset="UTF-8">
<title>Reporte Tickets — DIF El Márques</title>
<style>
  @page {{ size: A4 landscape; margin: 1.5cm 1.2cm; }}
  * {{ box-sizing: border-box; margin: 0; padding: 0; }}
  body {{ font-family: Arial, Helvetica, sans-serif; color: #1e293b;
          background: #fff; font-size: 11px; }}

  /* WATERMARK */
  .wm {{ position: fixed; top: 50%; left: 50%;
         transform: translate(-50%,-50%) rotate(-25deg);
         font-size: 90px; font-weight: 900; color: rgba(13,148,136,.04);
         text-transform: uppercase; letter-spacing: 10px;
         pointer-events: none; white-space: nowrap; z-index: 0; }}

  /* HEADER */
  .hdr {{ display: flex; align-items: center; justify-content: space-between;
          border-bottom: 3px solid #0d9488; padding-bottom: 12px; margin-bottom: 14px; }}
  .hdr-left {{ display: flex; align-items: center; gap: 14px; }}
  .hdr-logo {{ height: 52px; width: auto; object-fit: contain; }}
  .hdr-div  {{ width: 1px; height: 42px; background: #cbd5e1; }}
  .hdr-org  {{ font-size: 15px; font-weight: 800; color: #0d9488; line-height: 1.2; }}
  .hdr-sub  {{ font-size: 9px; color: #64748b; margin-top: 2px; }}
  .hdr-right {{ text-align: right; }}
  .hdr-title {{ font-size: 18px; font-weight: 800; color: #0f766e; }}
  .hdr-meta  {{ font-size: 9px; color: #64748b; margin-top: 3px; }}

  /* KPI */
  .kpi-row {{ display: flex; gap: 8px; margin-bottom: 12px; }}
  .kpi {{ flex: 1; border: 1px solid #e2e8f0; border-radius: 8px;
          padding: 10px 8px; text-align: center; background: #f8fafc; }}
  .kpi-num {{ font-size: 22px; font-weight: 800; line-height: 1; }}
  .kpi-lbl {{ font-size: 8.5px; color: #64748b; text-transform: uppercase;
              letter-spacing: .4px; margin-top: 3px; font-weight: 700; }}

  /* SECCIÓN */
  .sec {{ border: 1px solid #e2e8f0; border-radius: 8px;
          overflow: hidden; margin-bottom: 12px; }}
  .sec-hdr {{ background: #f8fafc; padding: 7px 12px; border-bottom: 1px solid #e2e8f0;
              font-size: 9.5px; font-weight: 800; color: #475569;
              text-transform: uppercase; letter-spacing: .5px;
              display: flex; align-items: center; justify-content: space-between; }}
  .sec-hdr span {{ color: #0d9488; }}

  /* TABLA */
  table {{ width: 100%; border-collapse: collapse; }}
  th {{ background: #1e293b; color: #94a3b8; font-size: 8.5px; font-weight: 700;
        text-transform: uppercase; letter-spacing: .4px; padding: 7px 9px;
        text-align: left; white-space: nowrap; }}
  td {{ padding: 6px 9px; border-bottom: 1px solid #f1f5f9; vertical-align: middle; }}
  tr:last-child td {{ border-bottom: none; }}

  /* FOOTER */
  .footer {{ border-top: 2px solid #e2e8f0; margin-top: 14px; padding-top: 8px;
             display: flex; justify-content: space-between; align-items: center; }}
  .footer-l {{ font-size: 8.5px; color: #94a3b8; }}
  .footer-c {{ font-size: 9px; font-weight: 700; color: #0d9488; }}
  .footer-r {{ font-size: 8.5px; color: #94a3b8; text-align: right; }}

  @media print {{
    body {{ -webkit-print-color-adjust: exact; print-color-adjust: exact; }}
  }}
</style>
</head>
<body>

<div class="wm">DIF</div>

<!-- HEADER -->
<div class="hdr">
  <div class="hdr-left">
    <img src="data:image/png;base64,{B64_LOGO}" class="hdr-logo" alt="IITIL Helpdesk">
    <div class="hdr-div"></div>
    <div>
      <div class="hdr-org">ITIL Helpdesk — DIF El Márques</div>
      <div class="hdr-sub">Sistema de Soporte Técnico · Departamento de Informática</div>
      <div class="hdr-sub">Generado: {fecha}</div>
    </div>
  </div>
  <div class="hdr-right">
    <div class="hdr-title">Reporte de Tickets</div>
    <div class="hdr-meta">{total} tickets en total</div>
  </div>
</div>

<!-- KPIs ESTADO -->
<div class="kpi-row">
  <div class="kpi">
    <div class="kpi-num" style="color:#0d9488">{total}</div>
    <div class="kpi-lbl">Total</div>
  </div>
  {kpi_html}
</div>

<!-- KPIs PRIORIDAD -->
<div style="margin-bottom:12px">
  <div style="font-size:9px;font-weight:700;color:#475569;text-transform:uppercase;
              letter-spacing:.4px;margin-bottom:6px">Distribución por Prioridad</div>
  <div class="kpi-row">{prio_html}</div>
</div>

<!-- TABLA TICKETS -->
<div class="sec">
  <div class="sec-hdr">
    Listado Completo de Tickets
    <span>{len(tickets)} registros</span>
  </div>
  <table>
    <thead>
      <tr>
        <th>Folio</th>
        <th>Título</th>
        <th>Usuario</th>
        <th>Área</th>
        <th>Estado</th>
        <th>Prioridad</th>
        <th>Técnico</th>
        <th>Fecha</th>
      </tr>
    </thead>
    <tbody>{rows_html}</tbody>
  </table>
</div>

<!-- FOOTER -->
<div class="footer">
  <div class="footer-l">
    ITIL Helpdesk — Sistema DIF de Soporte Técnico © 2026<br>
    Documento generado automáticamente
  </div>
  <div class="footer-c">DIF El Márques · Informática</div>
  <div class="footer-r">
    Reporte generado: {fecha}<br>
    Total de registros: {total}
  </div>
</div>

</body>
</html>"""

    @staticmethod
    def generar_reporte_word(stats, tickets):
        """Genera reporte .docx usando python-docx."""
        try:
            from docx import Document
            from docx.shared import Pt, RGBColor, Inches
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            import io
        except ImportError:
            raise ImportError("Instala python-docx: pip install python-docx")

        doc = Document()
        section = doc.sections[0]
        section.page_width  = Inches(11)
        section.page_height = Inches(8.5)
        section.left_margin = section.right_margin = Inches(1)

        title = doc.add_heading('Reporte de Tickets — DIF El Márques', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title.runs[0].font.color.rgb = RGBColor(0x0d, 0x94, 0x88)

        p = doc.add_paragraph(f'Generado el {datetime.now().strftime("%d/%m/%Y %H:%M")}')
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph()

        doc.add_heading('Resumen General', 1)
        p = doc.add_paragraph()
        run = p.add_run(f'Total de tickets: ')
        run.bold = True
        p.add_run(str(stats.get('total', 0)))

        for estado in stats.get('por_estado', []):
            p = doc.add_paragraph(style='List Bullet')
            run = p.add_run(f"{estado['nombre']}: ")
            run.bold = True
            p.add_run(str(estado['cantidad']))

        doc.add_paragraph()
        doc.add_heading('Listado de Tickets', 1)

        cols = ['Folio', 'Título', 'Usuario', 'Área', 'Estado', 'Prioridad', 'Técnico', 'Fecha']
        table = doc.add_table(rows=1, cols=len(cols))
        table.style = 'Table Grid'

        hdr = table.rows[0].cells
        for i, col in enumerate(cols):
            hdr[i].text = col
            run = hdr[i].paragraphs[0].runs[0]
            run.font.bold = True
            run.font.size = Pt(9)

        for t in tickets:
            row = table.add_row().cells
            vals = [
                t.get('folio',''), str(t.get('titulo',''))[:50],
                t.get('nombre_usuario',''), t.get('area_usuario',''),
                t.get('estado',''), t.get('prioridad',''),
                t.get('nombre_tecnico','') or '—', str(t.get('fecha_creacion',''))[:16],
            ]
            for i, val in enumerate(vals):
                row[i].text = val
                if row[i].paragraphs[0].runs:
                    row[i].paragraphs[0].runs[0].font.size = Pt(8)

        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()

    @staticmethod
    def calcular_metricas_desempeno(tickets):
        total     = len(tickets)
        resueltos = len([t for t in tickets if t.get('estado') in ('Resuelto', 'Cerrado')])
        return {
            'tickets_totales':    total,
            'tickets_resueltos':  resueltos,
            'tickets_pendientes': len([t for t in tickets if t.get('estado') == 'Abierto']),
            'tickets_en_proceso': len([t for t in tickets if t.get('estado') == 'En Proceso']),
            'tickets_criticos':   len([t for t in tickets if t.get('prioridad') == 'Critica']),
            'tasa_resolucion':    round((resueltos / total * 100), 2) if total else 0,
        }


# ── ALIAS para compatibilidad con user_routes.py ──────────────
# user_routes.py hace: from utils.reportes_pdf import generar_ticket_pdf
# Esta función genera el HTML del comprobante de UN ticket individual.

def generar_ticket_pdf(ticket, historial=None):
    """
    Genera HTML de comprobante para un ticket individual.
    Recibe el dict del ticket y opcionalmente el historial.
    """
    historial = historial or []
    fecha_gen = datetime.now().strftime('%d/%m/%Y %H:%M')

    color_est  = {'Abierto':'#3b82f6','En Proceso':'#f59e0b','Resuelto':'#22c55e','Cerrado':'#94a3b8'}
    color_prio = {'Baja':'#22c55e','Media':'#f59e0b','Alta':'#f97316','Crítica':'#ef4444','Critica':'#ef4444'}

    ce = color_est.get(ticket.get('estado',''), '#94a3b8')
    cp = color_prio.get(ticket.get('prioridad',''), '#94a3b8')

    hist_rows = ''
    for ev in historial:
        hist_rows += f"""<tr>
          <td style="font-size:10px;color:#64748b">{ev.get('fecha','')}</td>
          <td style="font-size:11px;font-weight:600">{ev.get('accion','')}</td>
          <td style="font-size:10px;color:#64748b">{ev.get('descripcion','') or '—'}</td>
          <td style="font-size:10px">{ev.get('nombre_completo','')}</td>
        </tr>"""

    solucion_html = ''
    if ticket.get('solucion'):
        solucion_html = f"""
        <div class="section">
          <div class="section-hdr" style="color:#065f46;background:#f0fdf4;border-color:#bbf7d0">
            ✅ Solución Aplicada
          </div>
          <div style="padding:14px">
            <p style="white-space:pre-wrap;font-size:12px;line-height:1.6">{ticket.get('solucion','')}</p>
            {'<p style="font-size:11px;color:#64748b;margin-top:8px">Observaciones: ' + ticket.get('observaciones_tecnico','') + '</p>' if ticket.get('observaciones_tecnico') else ''}
          </div>
        </div>"""

    return f"""<!DOCTYPE html>
<html lang="es">
<head>
<meta charset="UTF-8">
<title>Comprobante {ticket.get('folio','')}</title>
<style>
  @page {{ size: A4; margin: 1.5cm; }}
  *{{ box-sizing:border-box; margin:0; padding:0; }}
  body{{ font-family:'Segoe UI',Arial,sans-serif; color:#1e293b; background:#fff; font-size:12px; }}
  .no-print{{ background:#0d9488; color:#fff; padding:8px 16px; text-align:center;
    font-size:12px; font-weight:600; margin-bottom:14px; border-radius:8px;
    display:flex; align-items:center; justify-content:center; gap:10px; }}
  .print-btn{{ background:#fff; color:#0d9488; border:none; padding:5px 14px;
    border-radius:6px; font-weight:700; cursor:pointer; font-size:11px; }}
  .header{{ background:linear-gradient(135deg,#0d9488,#0f766e); color:#fff;
    padding:16px 20px; border-radius:10px; margin-bottom:14px;
    display:flex; justify-content:space-between; align-items:center; }}
  .header h1{{ font-size:16px; font-weight:800; margin-bottom:2px; }}
  .header p{{ font-size:10px; opacity:.8; }}
  .folio{{ font-family:monospace; font-size:18px; font-weight:800; }}
  .section{{ border:1px solid #e2e8f0; border-radius:9px; overflow:hidden; margin-bottom:12px; }}
  .section-hdr{{ background:#f8fafc; padding:8px 14px; border-bottom:1px solid #e2e8f0;
    font-size:10px; font-weight:700; color:#475569; text-transform:uppercase; letter-spacing:.4px; }}
  .grid{{ display:grid; grid-template-columns:1fr 1fr; gap:12px; padding:14px; }}
  .field label{{ font-size:9px; font-weight:700; color:#94a3b8; text-transform:uppercase;
    letter-spacing:.4px; display:block; margin-bottom:3px; }}
  .field .val{{ font-size:12px; color:#1e293b; font-weight:500; }}
  .badge{{ padding:3px 10px; border-radius:20px; font-size:10px;
    font-weight:700; color:#fff; display:inline-block; }}
  table{{ width:100%; border-collapse:collapse; }}
  th{{ background:#1e293b; color:#94a3b8; font-size:9px; font-weight:700;
    text-transform:uppercase; letter-spacing:.4px; padding:7px 10px;
    border-bottom:2px solid #334155; text-align:left; }}
  td{{ padding:7px 10px; border-bottom:1px solid #f1f5f9; vertical-align:middle; }}
  tr:last-child td{{ border-bottom:none; }}
  .footer{{ text-align:center; font-size:9px; color:#94a3b8; padding-top:10px;
    border-top:1px solid #e2e8f0; margin-top:12px; }}
  @media print{{
    .no-print{{ display:none !important; }}
    body{{ -webkit-print-color-adjust:exact; print-color-adjust:exact; }}
  }}
</style>
</head>
<body>
<div class="no-print">
  Comprobante de ticket · <strong>Ctrl+P</strong> para imprimir o guardar PDF
  <button class="print-btn" onclick="window.print()">🖨 Imprimir</button>
</div>

<div class="header">
  <div>
    <h1>🎫 Comprobante de Ticket</h1>
    <p>Sistema ITIL Helpdesk · DIF El Márques · Soporte Técnico</p>
  </div>
  <div class="folio">{ticket.get('folio','')}</div>
</div>

<div class="section">
  <div class="section-hdr">Información General</div>
  <div class="grid">
    <div class="field"><label>Estado</label><div class="val"><span class="badge" style="background:{ce}">{ticket.get('estado','')}</span></div></div>
    <div class="field"><label>Prioridad</label><div class="val"><span class="badge" style="background:{cp}">{ticket.get('prioridad','')}</span></div></div>
    <div class="field"><label>Usuario</label><div class="val">{ticket.get('nombre_usuario','')}</div></div>
    <div class="field"><label>Área</label><div class="val">{ticket.get('area_usuario','')}</div></div>
    <div class="field"><label>Tipo de Problema</label><div class="val">{ticket.get('tipo_problema','') or '—'}</div></div>
    <div class="field"><label>Técnico Asignado</label><div class="val">{ticket.get('nombre_tecnico','') or 'Sin asignar'}</div></div>
    <div class="field"><label>Fecha de Creación</label><div class="val">{str(ticket.get('fecha_creacion',''))[:16]}</div></div>
    <div class="field"><label>Fecha de Resolución</label><div class="val">{str(ticket.get('fecha_resolucion','') or '—')[:16]}</div></div>
  </div>
</div>

<div class="section">
  <div class="section-hdr">Descripción del Problema</div>
  <div style="padding:14px">
    <p style="font-weight:700;margin-bottom:6px;font-size:13px">{ticket.get('titulo','')}</p>
    <p style="white-space:pre-wrap;font-size:12px;color:#334155;line-height:1.6">{ticket.get('descripcion','')}</p>
  </div>
</div>

{solucion_html}

{'<div class="section"><div class="section-hdr">Historial de Cambios</div><table><thead><tr><th>Fecha</th><th>Acción</th><th>Descripción</th><th>Responsable</th></tr></thead><tbody>' + hist_rows + '</tbody></table></div>' if historial else ''}

<div class="footer">
  ITIL Helpdesk — Sistema DIF de Soporte Técnico © 2026 · Comprobante generado: {fecha_gen}
</div>
</body></html>"""