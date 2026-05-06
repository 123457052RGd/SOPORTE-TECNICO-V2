from docx import Document
from docx.shared import Inches
import os

def generar_ticket_word(ticket, historial, ruta_salida):
    doc = Document()

    # Logos
    logo_dif = "static/img/logodif.png"
    logo_sys = "static/img/logo.png"

    table = doc.add_table(rows=1, cols=3)
    row = table.rows[0]

    if os.path.exists(logo_dif):
        row.cells[0].paragraphs[0].add_run().add_picture(logo_dif, width=Inches(1.2))

    row.cells[1].text = "SISTEMA DE SOPORTE TÉCNICO"

    if os.path.exists(logo_sys):
        row.cells[2].paragraphs[0].add_run().add_picture(logo_sys, width=Inches(1.2))

    doc.add_heading('Reporte de Ticket', level=1)

    doc.add_paragraph(f"Folio: {ticket.get('folio')}")
    doc.add_paragraph(f"Usuario: {ticket.get('nombre_usuario')}")
    doc.add_paragraph(f"Área: {ticket.get('area_usuario')}")
    doc.add_paragraph(f"Estado: {ticket.get('estado')}")
    doc.add_paragraph(f"Prioridad: {ticket.get('prioridad')}")

    doc.add_heading('Descripción', level=2)
    doc.add_paragraph(ticket.get('descripcion', ''))

    doc.add_heading('Historial', level=2)

    for h in historial:
        doc.add_paragraph(
            f"{h.get('fecha')} - {h.get('accion')} - {h.get('descripcion')}"
        )

    doc.save(ruta_salida)